import argparse
import os
import platform
from io import BytesIO

from docx import Document
from docx.shared import Cm

from properties import properties

SYSTEM = platform.system()


def get_image(x, y):
    if SYSTEM == "Linux":
        is_wayland = os.environ.get("WAYLAND_DISPLAY")
        is_x11 = os.environ.get("DISPLAY")
        if is_wayland:
            # raise Exception("Wayland is not yet implemented")
            from screenshots.pil_screenshots import pil_screenshot
            screenshot = pil_screenshot()
            x = 0
        elif is_x11:
            from screenshots.xlib_screenshots import xlib_screenshot
            im, relative_x, relative_y = xlib_screenshot()
            x -= relative_x
            y -= relative_y
        else:
            raise Exception("Unsupported display server")
    elif SYSTEM == "Windows":
        from screenshots.pyautogui_screenshots import pyautogui_screenshot
        im, relative_x, relative_y = pyautogui_screenshot()
        x -= relative_x
        y -= relative_y
    else:
        raise Exception("Unsupported operating system: {}".format(SYSTEM))
    return im, x, y


def write_to_file(im):
    im_bytes = BytesIO()
    im.save(im_bytes, format="png")
    im_bytes.seek(0)
    if os.path.exists(properties.file_path):
        doc = Document(properties.file_path)
    else:
        os.makedirs(os.path.dirname(properties.file_path))
        doc = Document()
    page_width = doc.sections[0].page_width.cm - doc.sections[0].left_margin.cm - doc.sections[0].right_margin.cm
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(im_bytes, width=Cm(page_width))
    doc.save(properties.file_path)


if __name__ == "__main__":

    parser = argparse.ArgumentParser()
    parser.add_argument("-f", "--file", type=str, required=True, help="target file to store the screenshots")
    parser.add_argument("-m", "--marker", type=str, choices=["arrow", "rect", "circle", "bullet", "none"],
                        help="Symbol that marks the clicked area. Default is 'none' ('none' is providing no marker at all)")
    args = parser.parse_args()

    if args.marker:
        properties.marker = args.marker

    properties.file_path = args.file

    try:
        print("Use middle click to take a screenshot!")
        if SYSTEM == "Windows":
            from input_listeners.windows import attach_listener as win_al
            win_al()
        elif SYSTEM == "Linux":
            from input_listeners.linux import attach_listener as lin_al
            lin_al()
        else:
            raise Exception("Unsupported operating system: {}".format(SYSTEM))
    except KeyboardInterrupt:
        print("\n Good bye!")
